from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
from typing import List, Optional, Dict
import os
import json
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
import pdfplumber
import tempfile
from app.core.config import settings

router = APIRouter()

class ATSScore(BaseModel):
    overall_score: int
    content_score: int
    format_score: int
    keyword_score: int

class ResumeAnalysis(BaseModel):
    ats_score: ATSScore
    skills_match: Dict[str, bool]
    missing_skills: List[str]
    suggestions: List[str]
    summary: str
    extracted_text: Optional[str] = None

class JobMatchRequest(BaseModel):
    resume_text: str
    job_description: str

class JobMatchResponse(BaseModel):
    match_percentage: int
    matched_skills: List[str]
    missing_skills: List[str]
    suggestions: List[str]

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def sanitize_analysis_result(result: dict) -> dict:
    import re
    sanitized = {}
    
    ats_score = result.get("ats_score", {})
    if not isinstance(ats_score, dict):
        ats_score = {}
        
    def coerce_to_int(val, default=0) -> int:
        if isinstance(val, (int, float)):
            return int(val)
        if isinstance(val, str):
            val = val.split('/')[0].replace('%', '').strip()
            try:
                return int(val)
            except ValueError:
                match = re.search(r'\d+', val)
                if match:
                    return int(match.group())
        return default

    sanitized_score = {
        "overall_score": coerce_to_int(ats_score.get("overall_score"), 50),
        "content_score": coerce_to_int(ats_score.get("content_score"), 50),
        "format_score": coerce_to_int(ats_score.get("format_score"), 50),
        "keyword_score": coerce_to_int(ats_score.get("keyword_score"), 50)
    }
    sanitized["ats_score"] = sanitized_score

    skills_match = result.get("skills_match", {})
    sanitized_skills = {}
    if isinstance(skills_match, dict):
        for k, v in skills_match.items():
            if isinstance(v, bool):
                sanitized_skills[str(k)] = v
            elif isinstance(v, str):
                val_lower = v.lower()
                sanitized_skills[str(k)] = val_lower in ("true", "yes", "found", "y", "1")
            elif isinstance(v, (int, float)):
                sanitized_skills[str(k)] = bool(v)
            else:
                sanitized_skills[str(k)] = False
    elif isinstance(skills_match, list):
        for item in skills_match:
            sanitized_skills[str(item)] = True
    sanitized["skills_match"] = sanitized_skills

    missing = result.get("missing_skills", [])
    if isinstance(missing, list):
        sanitized["missing_skills"] = [str(x) for x in missing]
    else:
        sanitized["missing_skills"] = []

    suggestions = result.get("suggestions", [])
    if isinstance(suggestions, list):
        sanitized["suggestions"] = [str(x) for x in suggestions]
    else:
        sanitized["suggestions"] = []

    summary = result.get("summary", "")
    sanitized["summary"] = str(summary) if summary else "No summary provided."
    
    return sanitized

def sanitize_job_match(result: dict) -> dict:
    import re
    sanitized = {}
    
    def coerce_to_int(val, default=0) -> int:
        if isinstance(val, (int, float)):
            return int(val)
        if isinstance(val, str):
            val = val.split('/')[0].replace('%', '').strip()
            try:
                return int(val)
            except ValueError:
                match = re.search(r'\d+', val)
                if match:
                    return int(match.group())
        return default
        
    sanitized["match_percentage"] = coerce_to_int(result.get("match_percentage"), 0)
    
    matched = result.get("matched_skills", [])
    sanitized["matched_skills"] = [str(x) for x in matched] if isinstance(matched, list) else []
    
    missing = result.get("missing_skills", [])
    sanitized["missing_skills"] = [str(x) for x in missing] if isinstance(missing, list) else []
    
    suggestions = result.get("suggestions", [])
    sanitized["suggestions"] = [str(x) for x in suggestions] if isinstance(suggestions, list) else []
    
    return sanitized

@router.post("/analyze", response_model=ResumeAnalysis)
async def analyze_resume(file: UploadFile = File(...)):
    contents = await file.read()
    
    # Use a secure, cross-platform temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(contents)
        tmp_file_path = tmp_file.name
    
    try:
        resume_text = extract_text_from_pdf(tmp_file_path)
    except Exception:
        resume_text = contents.decode('utf-8', errors='ignore')
    finally:
        try:
            if os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)
        except Exception:
            pass
    
    parser = JsonOutputParser()
    
    prompt = ChatPromptTemplate.from_template("""
    Analyze this resume and provide ATS evaluation:
    
    Resume: {resume_text}
    
    Provide JSON with:
    - ats_score: {{overall_score, content_score, format_score, keyword_score}}
    - skills_match: {{skill: found}}
    - missing_skills: []
    - suggestions: []
    - summary: "brief analysis"
    
    Focus on: keywords, formatting, action verbs, quantifiable achievements
    """)
    
    models = []
    for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
        if m and m not in models:
            models.append(m)
            
    last_error = None
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.3,
                max_retries=1
            )
            chain = prompt | llm | parser
            result = chain.invoke({"resume_text": resume_text[:3000]})
            
            # Sanitize result to match Pydantic schema types and structure
            sanitized = sanitize_analysis_result(result)
            sanitized["extracted_text"] = resume_text
            return ResumeAnalysis(**sanitized)
        except Exception as e:
            last_error = e
            import sys
            print(f"Failed resume analysis using model {model}: {e}", file=sys.stderr)
            continue
            
    import traceback
    import sys
    import re
    print("ERROR during resume analysis (all models failed). Falling back to heuristic analysis...", file=sys.stderr)
    
    # Fallback heuristic analysis
    has_email = bool(re.search(r'([a-zA-Z0-9._%-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,6})', resume_text))
    has_phone = bool(re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text))
    
    content_score = 75
    format_score = 80
    
    # Check for metrics
    has_metrics = bool(re.search(r'[\d]+%|[\d]+\s*percent|\$[\d]+[kKmM]?|[\d]+\s*(?:million|billion|x|X|times)', resume_text))
    if has_metrics:
        content_score += 15
    else:
        content_score -= 10
        
    # Check for action verbs
    action_verbs = ["lead", "led", "manage", "managed", "develop", "developed", "design", "designed",
                    "build", "built", "create", "created", "optimize", "optimized", "improve", "improved",
                    "implement", "implemented"]
    has_verbs = any(re.search(rf'\b{verb}\b', resume_text, re.IGNORECASE) for verb in action_verbs)
    if has_verbs:
        content_score += 5
    else:
        content_score -= 5
        
    if has_email:
        format_score += 10
    else:
        format_score -= 20
        
    if has_phone:
        format_score += 5
    else:
        format_score -= 10
        
    # Skills matching
    tech_skills = ["react", "python", "java", "javascript", "typescript", "node.js", "html", "css", 
                   "sql", "docker", "aws", "git", "kubernetes", "c++", "go", "angular", "vue", "django"]
    found_skills = {}
    missing_skills = []
    
    for skill in tech_skills:
        pattern = rf'\b{re.escape(skill)}\b'
        if skill == "node.js":
            pattern = r'\bnode(?:\.js)?\b'
        elif skill == "c++":
            pattern = r'\bc\+\+\b'
        
        if re.search(pattern, resume_text, re.IGNORECASE):
            pretty_name = skill.title() if skill not in ["sql", "html", "css", "aws"] else skill.upper()
            if skill == "javascript": pretty_name = "JavaScript"
            elif skill == "typescript": pretty_name = "TypeScript"
            elif skill == "node.js": pretty_name = "Node.js"
            found_skills[pretty_name] = True
        else:
            pretty_name = skill.title() if skill not in ["sql", "html", "css", "aws"] else skill.upper()
            if skill == "javascript": pretty_name = "JavaScript"
            elif skill == "typescript": pretty_name = "TypeScript"
            elif skill == "node.js": pretty_name = "Node.js"
            missing_skills.append(pretty_name)
            
    found_count = len(found_skills)
    keyword_score = min(100, 50 + (found_count * 5))
    
    overall_score = int((content_score + format_score + keyword_score) / 3)
    
    suggestions = []
    if not has_email:
        suggestions.append("Critical: Add a professional email address to the header.")
    if not has_phone:
        suggestions.append("Warning: Add your contact number so recruiters can easily reach you.")
    if not has_metrics:
        suggestions.append("Improvement: Add quantifiable metrics and percentages (e.g., 'increased sales by 15%') to prove your impact.")
    if not has_verbs:
        suggestions.append("Improvement: Start your experience bullet points with strong action verbs (e.g., 'Delivered', 'Spearheaded').")
    if found_count < 4:
        suggestions.append("Improvement: Incorporate more core technical skills related to your target role.")
    else:
        suggestions.append("Good: Standard section structures and keywords are well-balanced.")
        
    suggestions.append("Action: Customize your resume summary to match specific keywords from your target job description.")
    
    summary = "Heuristic Audit (API Fallback): "
    if found_count > 0:
        skills_str = ", ".join(list(found_skills.keys())[:5])
        summary += f"Resume highlights skills in {skills_str}. "
    else:
        summary += "Resume parsed with general details. "
        
    if overall_score >= 80:
        summary += "Overall structure and keyword representation are excellent. Ready for application."
    elif overall_score >= 60:
        summary += "Strong foundation detected, but some key additions (like metrics or formatting improvements) will optimize it for ATS screening."
    else:
        summary += "Critical adjustments are needed to formatting and contact details before submitting to job boards."
        
    result = {
        "ats_score": {
            "overall_score": overall_score,
            "content_score": content_score,
            "format_score": format_score,
            "keyword_score": keyword_score
        },
        "skills_match": found_skills,
        "missing_skills": missing_skills[:4],
        "suggestions": suggestions,
        "summary": summary,
        "extracted_text": resume_text
    }
    
    return ResumeAnalysis(**result)

@router.post("/match-job", response_model=JobMatchResponse)
async def match_resume_to_job(request: JobMatchRequest):
    prompt = ChatPromptTemplate.from_template("""
    Compare resume against job description and calculate match:
    
    Resume: {resume_text}
    Job Description: {job_description}
    
    Return JSON: {{match_percentage, matched_skills, missing_skills, suggestions}}
    """)
    
    models = []
    for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
        if m and m not in models:
            models.append(m)
            
    last_error = None
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.3,
                max_retries=1
            )
            chain = prompt | llm | JsonOutputParser()
            result = chain.invoke(request.dict())
            
            # Sanitize result to match Pydantic schema types and structure
            sanitized = sanitize_job_match(result)
            return JobMatchResponse(**sanitized)
        except Exception as e:
            last_error = e
            import sys
            print(f"Failed resume matching using model {model}: {e}", file=sys.stderr)
            continue
            
    import traceback
    import sys
    import re
    print("ERROR during resume matching (all models failed). Falling back to heuristic match...", file=sys.stderr)
    
    # Simple keyword match ratio
    resume_text_lower = request.resume_text.lower()
    jd_text_lower = request.job_description.lower()
    
    # List of technical terms/skills we care about
    common_skills = [
        "react", "python", "java", "javascript", "typescript", "node", "html", "css", 
        "sql", "docker", "aws", "git", "kubernetes", "c++", "go", "angular", "vue", 
        "django", "spring", "flask", "postgresql", "mongodb", "mysql", "redis",
        "rest api", "graphql", "devops", "ci/cd", "agile", "scrum", "junit", "maven"
    ]
    
    matched_skills = []
    missing_skills = []
    
    for skill in common_skills:
        # Check if skill is in job description
        skill_in_jd = False
        if skill == "c++":
            skill_in_jd = r"c\+\+" in jd_text_lower
        else:
            skill_in_jd = rf"\b{re.escape(skill)}\b" in jd_text_lower
            
        if skill_in_jd:
            # Check if skill is in resume
            skill_in_resume = False
            if skill == "c++":
                skill_in_resume = r"c\+\+" in resume_text_lower
            else:
                skill_in_resume = rf"\b{re.escape(skill)}\b" in resume_text_lower
                
            pretty = skill.upper() if skill in ["sql", "html", "css", "aws", "api", "ci/cd"] else skill.title()
            if pretty == "Rest Api": pretty = "REST API"
            elif pretty == "Postgresql": pretty = "PostgreSQL"
            elif pretty == "Mongodb": pretty = "MongoDB"
            
            if skill_in_resume:
                matched_skills.append(pretty)
            else:
                missing_skills.append(pretty)
                
    total_jd_skills = len(matched_skills) + len(missing_skills)
    if total_jd_skills > 0:
        match_percentage = int((len(matched_skills) / total_jd_skills) * 100)
    else:
        # Default fallback match based on word overlap
        match_percentage = 65
        
    # Bound match percentage reasonably
    match_percentage = max(30, min(95, match_percentage))
    
    suggestions = []
    if missing_skills:
        suggestions.append(f"Add target keywords from job description: {', '.join(missing_skills[:3])}")
    suggestions.append("Tailor your professional summary to explicitly match the target job title.")
    suggestions.append("Describe project work using the same methodologies (e.g. Agile, DevOps) listed in the posting.")
    
    return JobMatchResponse(
        match_percentage=match_percentage,
        matched_skills=matched_skills,
        missing_skills=missing_skills[:5],
        suggestions=suggestions
    )

class RewriteRequest(BaseModel):
    section: Optional[str] = ""
    style: str = "professional"
    job_description: Optional[str] = None
    mode: str = "improve"  # "improve" or "scratch"

class ChatResumeRequest(BaseModel):
    resume_text: str
    question: str

@router.post("/rewrite")
async def rewrite_resume_section(request: RewriteRequest):
    models = []
    for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
        if m and m not in models:
            models.append(m)
            
    last_error = None
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.7,
                max_retries=1
            )
            
            if request.mode == "scratch":
                system_prompt = (
                    "You are an expert resume writer. Generate a comprehensive, professional, and ATS-friendly resume from scratch.\n\n"
                    "Job Description (Target Role):\n"
                    "{job_description}\n\n"
                    "Guidelines:\n"
                    "1. Tailor the resume specifically to the job description provided above, highlighting relevant skills and keywords.\n"
                    "2. Structure it with clear sections: Contact Info, Professional Summary, Work Experience (with bullet points using action verbs and placeholders for metrics like '[X]%'), Education, and Skills.\n"
                    "3. Use a clean, plain text markdown format."
                )
                prompt = ChatPromptTemplate.from_messages([
                    ("system", system_prompt),
                    ("human", "Generate a new resume based on the target job description.")
                ])
                chain = prompt | llm | StrOutputParser()
                result = chain.invoke({"job_description": request.job_description or "General Professional Resume Outline"})
            else:
                # mode == "improve"
                system_prompt = (
                    "You are an expert resume rewriter. Rewrite the following resume section to make it highly impactful, professional, and ATS-friendly.\n\n"
                    "Original Section:\n"
                    "{section}\n\n"
                )
                if request.job_description:
                    system_prompt += (
                        "Job Description (Target Role):\n"
                        "{job_description}\n\n"
                        "Strict Instructions:\n"
                        "Align the rewritten text specifically to match the requirements, skills, and key phrases found in the job description.\n\n"
                    )
                system_prompt += (
                    "Guidelines:\n"
                    "1. Improve action verbs, sentence structure, and make bullet points punchy.\n"
                    "2. Whenever possible, structure sentences to show achievement and impact (e.g. Accomplished [X] as measured by [Y] by doing [Z]).\n"
                    "3. Retain the core factual details from the original section but present them with maximum impact."
                )
                prompt = ChatPromptTemplate.from_messages([
                    ("system", system_prompt),
                    ("human", "Rewrite the resume section in a {style} style.")
                ])
                chain = prompt | llm | StrOutputParser()
                result = chain.invoke({
                    "section": request.section,
                    "style": request.style,
                    "job_description": request.job_description or ""
                })
                
            return {"rewritten": result}
        except Exception as e:
            last_error = e
            import sys
            print(f"Failed resume rewrite using model {model}: {e}", file=sys.stderr)
            continue
            
    import traceback
    import sys
    import re
    print("ERROR during resume rewriting (all models failed). Falling back to heuristic rewrite...", file=sys.stderr)
    
    # Heuristic rewriter
    if request.mode == "scratch":
        # Generate a high-quality formatted markdown resume template
        role = request.job_description or "Software Engineer"
        result = f"""# [YOUR NAME]
[City, Country] | [Phone Number] | [Email Address] | [LinkedIn Link] | [GitHub Link]

## PROFESSIONAL SUMMARY
Results-driven Professional with experience in developing, optimizing, and scaling web applications. Proven track record of leveraging industry best practices and modern tech stacks to deliver high-performing solutions that meet business requirements.

## WORK EXPERIENCE
**Senior Software Engineer** | [Company Name] | [Start Date] – Present
- Spearheaded the design and development of core application services, resulting in a **25% improvement in processing latency**.
- Collaborated with cross-functional teams to integrate new features, reducing time-to-market by **15%**.
- Mentored junior engineers and advocated for clean code practices, improving code review turnaround time by **30%**.

**Software Developer** | [Previous Company Name] | [Start Date] – [End Date]
- Implemented responsive user interface components using React and styled components, enhancing user engagement metrics by **12%**.
- Optimized database queries, resolving critical performance bottlenecks and decreasing system CPU usage by **18%**.

## EDUCATION
**Bachelor of Science in Computer Science** | [University Name] | [Graduation Year]

## TECHNICAL SKILLS
- **Languages:** JavaScript, TypeScript, Python, Java, SQL, HTML, CSS
- **Frameworks & Libs:** React, Node.js, Express, Spring Boot
- **Tools & Platforms:** Git, Docker, AWS, PostgreSQL, MongoDB, Linux
"""
    else:
        # mode == "improve"
        original = request.section or ""
        
        # Replace weak verbs with strong action verbs
        replacements = {
            r"\b(?:I|we)\s+worked\s+on\b": "Spearheaded design and development of",
            r"\b(?:I\s+was|we\s+were)\s+responsible\s+for\b": "Coordinated and executed deployment of",
            r"\b(?:I|we)\s+helped\s+with\b": "Collaborated with cross-functional teams to deliver",
            r"\b(?:I|we)\s+made\b": "Architected and delivered",
            r"\b(?:I|we)\s+built\b": "Developed and integrated",
            r"\b(?:I|we)\s+fixed\b": "Troubleshot and resolved",
            r"\b(?:I|we)\s+managed\b": "Orchestrated and managed",
            r"\b(?:I|we)\s+used\b": "Leveraged",
            r"\b(?:I|we)\s+created\b": "Formulated and launched",
            r"\b(?:I|we)\s+wrote\b": "Authored and tested",
            r"\b(?:I|we)\s+improved\b": "Optimized and enhanced",
            r"\bworked\b": "implemented key processes",
            r"\bhelped\b": "supported system scalability",
            r"\bbuilt\b": "developed",
            r"\bfixed\b": "resolved"
        }
        
        improved = original
        for pattern, replacement in replacements.items():
            improved = re.sub(pattern, replacement, improved, flags=re.IGNORECASE)
            
        # Ensure bullet points start with uppercase and look clean
        lines = improved.split('\n')
        cleaned_lines = []
        for line in lines:
            line_str = line.strip()
            if line_str.startswith('-') or line_str.startswith('*'):
                # Capitalize first character after bullet marker
                bullet_marker = line_str[0]
                rest = line_str[1:].strip()
                if rest:
                    line_str = f"{bullet_marker} {rest[0].upper()}{rest[1:]}"
            cleaned_lines.append(line_str)
        improved = "\n".join(cleaned_lines)
        
        # Add metrics suggestion if no metrics found
        if not re.search(r'[\d]+%|[\d]+\s*percent|\$[\d]+[kKmM]?', improved):
            improved += "\n- [Metric Placeholder] Increased operational efficiency by **20%** through refactoring and optimization."
            
        result = improved
        
    return {"rewritten": result}
        
    return {"rewritten": result}

@router.post("/chat")
async def chat_resume(request: ChatResumeRequest):
    prompt = ChatPromptTemplate.from_template("""
    You are an AI Resume Assistant. Help the user improve their resume.
    
    Resume Context: {resume_text}
    
    User Question: {question}
    """)
    
    models = []
    for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
        if m and m not in models:
            models.append(m)
            
    last_error = None
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.3,
                max_retries=1
            )
            chain = prompt | llm | StrOutputParser()
            result = chain.invoke({"resume_text": request.resume_text, "question": request.question})
            return {"answer": result}
        except Exception as e:
            last_error = e
            import sys
            print(f"Failed resume chat using model {model}: {e}", file=sys.stderr)
            continue
            
    import traceback
    import sys
    import re
    print("ERROR during resume chat (all models failed). Falling back to heuristic chat...", file=sys.stderr)
    
    # Heuristic chatbot response
    question = request.question.lower()
    resume_text = request.resume_text
    
    response = ""
    if "skills" in question or "skill" in question:
        # Match potential skills from resume
        tech_skills = ["react", "python", "java", "javascript", "typescript", "node.js", "html", "css", 
                       "sql", "docker", "aws", "git", "kubernetes", "c++", "go", "angular", "vue", "django"]
        found = [s.title() for s in tech_skills if re.search(rf"\b{re.escape(s)}\b", resume_text, re.IGNORECASE)]
        if found:
            response = f"Based on your resume, I detected the following core skills: {', '.join(found)}. To optimize your resume for ATS, consider grouping them clearly by category (e.g. Languages, Frameworks, Developer Tools) and adding any relevant missing skills from target job descriptions."
        else:
            response = "I couldn't identify any standard technical skills in the plain text. I suggest adding a dedicated 'Skills' section listing your programming languages, libraries, and developer tools to improve parser matches."
    elif "experience" in question or "work" in question or "job" in question:
        response = "Your experience section forms the core of your ATS scan. I recommend writing bullet points in the STAR format: **Situation, Task, Action, Result**. Always begin each point with a strong action verb (e.g. 'Created', 'Designed', 'Optimized') and try to quantify your achievements with clear numbers or percentages (e.g. 'improved system throughput by 20%')."
    elif "metric" in question or "percent" in question or "number" in question:
        response = "Recruiters and ATS algorithms love quantified impact! Look for places in your resume where you can answer: How much? How many? How fast? For example, instead of saying 'Wrote SQL queries to fetch data', say 'Optimized database search queries, reducing response times by 35% and improving concurrent user support.'"
    elif "format" in question or "design" in question or "template" in question:
        response = "For optimal ATS scanning, keep formatting simple and standard. Avoid multi-column layouts, graphics, icons, progress bars, or tables inside text. Use standard fonts (like Arial or Calibri) and clear headings (such as 'Work Experience' and 'Education'). You can download your resume in our high-quality Word template using the button in the top right!"
    else:
        response = f"Thank you for your question about: '{request.question}'. As an automated fallback, here is an important tip: make sure your resume contains a clean heading structure, lists relevant target keywords, starts bullet points with strong action verbs, and quantifies your achievements. You can also rewrite specific sections using our 'Resume Rewriter' tab."
        
    return {"answer": response}

class DocxConversionRequest(BaseModel):
    text: str
    mode: str = "improve"

@router.post("/convert-docx")
async def convert_resume_to_docx(request: DocxConversionRequest):
    try:
        import io
        import base64
        import re
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        def add_p_border_bottom(paragraph, color_hex="3b82f6", size="12"):
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), size)
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), color_hex)
            pBdr.append(bottom)
            
        doc = Document()
        for s in doc.sections:
            s.top_margin = Inches(0.75)
            s.bottom_margin = Inches(0.75)
            s.left_margin = Inches(0.75)
            s.right_margin = Inches(0.75)
            
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        lines = [line.strip() for line in request.text.split('\n')]
        
        email_regex = re.compile(r'([a-zA-Z0-9._%-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,6})')
        phone_regex = re.compile(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        linkedin_regex = re.compile(r'(linkedin\.com/in/[a-zA-Z0-9_-]+)', re.IGNORECASE)
        github_regex = re.compile(r'(github\.com/[a-zA-Z0-9_-]+)', re.IGNORECASE)
        
        name = ""
        email = ""
        phone = ""
        linkedin = ""
        github = ""
        
        content_lines = []
        first_non_empty = ""
        for line in lines:
            if not line:
                continue
            
            email_match = email_regex.search(line)
            if email_match and not email:
                email = email_match.group(1)
                
            phone_match = phone_regex.search(line)
            if phone_match and not phone:
                phone = phone_match.group(0)
                
            linkedin_match = linkedin_regex.search(line)
            if linkedin_match and not linkedin:
                linkedin = linkedin_match.group(1)
                
            github_match = github_regex.search(line)
            if github_match and not github:
                github = github_match.group(1)
                
            is_contact = (
                '@' in line or
                'linkedin.com' in line.lower() or
                'github.com' in line.lower() or
                (phone and phone in line) or
                (line == '|')
            )
            if is_contact:
                continue
                
            if not first_non_empty and not line.startswith('#'):
                first_non_empty = line
                
            content_lines.append(line)
            
        first_line = next((l for l in lines if l), "")
        if first_line.startswith('#'):
            name = first_line.lstrip('#').strip()
        elif first_non_empty and len(first_non_empty) < 40 and ':' not in first_non_empty:
            name = first_non_empty
            if first_non_empty in content_lines:
                content_lines.remove(first_non_empty)
        else:
            name = "Resume"
            
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(name)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        
        contact_parts = []
        if email:
            contact_parts.append(email)
        if phone:
            contact_parts.append(phone)
        if linkedin:
            contact_parts.append(linkedin)
        if github:
            contact_parts.append(github)
            
        if contact_parts:
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_p.paragraph_format.space_before = Pt(2)
            contact_p.paragraph_format.space_after = Pt(12)
            
            contact_text = "  |  ".join(contact_parts)
            run = contact_p.add_run(contact_text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            add_p_border_bottom(contact_p, color_hex="3b82f6", size="12")
            
        sections_list = []
        current_sec = None
        
        def is_header(l):
            if l.startswith('#'):
                return True
            upper = l.upper()
            common = [
                "SUMMARY", "PROFESSIONAL SUMMARY", "EXPERIENCE", "WORK EXPERIENCE",
                "EDUCATION", "SKILLS", "TECHNICAL SKILLS", "PROJECTS", "CERTIFICATIONS",
                "ORGANIZATIONS", "LANGUAGES", "COURSES", "INTERNSHIPS", "SUMMARY OF QUALIFICATIONS"
            ]
            if upper in common:
                return True
            if len(l) > 2 and len(l) < 35 and upper == l and '|' not in l and '*' not in l and ':' not in l and '-' not in l:
                return True
            return False

        for line in content_lines:
            if is_header(line):
                if current_sec:
                    sections_list.append(current_sec)
                current_sec = {
                    "title": line.lstrip('#').strip("-: "),
                    "items": []
                }
            else:
                if not current_sec:
                    current_sec = {"title": "Profile", "items": []}
                current_sec["items"].append(line)
                
        if current_sec:
            sections_list.append(current_sec)
            
        for sec in sections_list:
            if not sec["items"]:
                continue
                
            head_p = doc.add_paragraph()
            head_p.paragraph_format.space_before = Pt(14)
            head_p.paragraph_format.space_after = Pt(6)
            head_p.paragraph_format.keep_with_next = True
            
            run = head_p.add_run(sec["title"].upper())
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
            
            add_p_border_bottom(head_p, color_hex="cbd5e1", size="6")
            
            for item in sec["items"]:
                is_bullet = item.startswith('-') or item.startswith('*') or item.startswith('•') or item.startswith('+')
                item_clean = re.sub(r'^[\-\*\•\+]\s*', '', item)
                
                parts = re.split(r'(\*\*.*?\*\*)', item_clean)
                
                if is_bullet:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    
                    if '|' in item or '  ' in item or (' - ' in item and len(item) < 100):
                        header_parts = [hp.strip() for hp in re.split(r'[|]|\s{2,}', item) if hp.strip()]
                        if len(header_parts) >= 2:
                            p.paragraph_format.space_before = Pt(6)
                            p.paragraph_format.space_after = Pt(2)
                            
                            table = doc.add_table(rows=1, cols=2)
                            table.autofit = False
                            table.allow_autofit = False
                            table.columns[0].width = Inches(5.0)
                            table.columns[1].width = Inches(2.0)
                            
                            row = table.rows[0]
                            
                            left_cell_p = row.cells[0].paragraphs[0]
                            left_cell_p.paragraph_format.space_after = Pt(0)
                            
                            run_title = left_cell_p.add_run(header_parts[0])
                            run_title.font.bold = True
                            run_title.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
                            
                            if len(header_parts) > 2:
                                left_cell_p.add_run("  •  ")
                                run_sub = left_cell_p.add_run(", ".join(header_parts[1:-1]))
                                run_sub.font.italic = True
                                run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                            
                            right_cell_p = row.cells[1].paragraphs[0]
                            right_cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            right_cell_p.paragraph_format.space_after = Pt(0)
                            
                            run_date = right_cell_p.add_run(header_parts[-1])
                            run_date.font.bold = True
                            run_date.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                            
                            pPr = p._element.getparent()
                            if pPr is not None:
                                pPr.remove(p._element)
                            continue
                            
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        clean_part = part[2:-2]
                        r = p.add_run(clean_part)
                        r.font.bold = True
                    else:
                        r = p.add_run(part)
                        
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        
        return {"word_base64": base64.b64encode(docx_bytes).decode()}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")

class CoverLetterRequest(BaseModel):
    resume_text: str
    job_description: str
    company_name: Optional[str] = None
    role: Optional[str] = None

@router.post("/generate-cover-letter")
async def generate_cover_letter(request: CoverLetterRequest):
    prompt = ChatPromptTemplate.from_template("""
    You are an expert career coach. Write a highly tailored, professional, and compelling cover letter for a candidate based on their resume and the target job description.
    
    Candidate Resume Details:
    {resume_text}
    
    Target Job Details:
    Role: {role}
    Company: {company_name}
    Job Description:
    {job_description}
    
    Write a 350-450 word cover letter that highlights the most relevant skills and projects from the resume matching the job requirements. Keep it professional, engaging, and clear. Do not use generic placeholders like [Date] or [Insert Name] at the top; write a complete, drop-in ready letter. Use the candidate's name from their resume if found, otherwise use a placeholder name.
    """)
    
    models = [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "openrouter/free"]
    last_error = None
    
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.7,
                max_retries=1
            )
            chain = prompt | llm | StrOutputParser()
            result = await chain.ainvoke({
                "resume_text": request.resume_text,
                "job_description": request.job_description,
                "role": request.role or "Software Engineer",
                "company_name": request.company_name or "Target Company"
            })
            return {"cover_letter": result}
        except Exception as e:
            last_error = e
            continue
            
    return {"cover_letter": "Dear Hiring Manager,\n\nI am writing to express my strong interest in the " + (request.role or "Software Engineer") + " position at " + (request.company_name or "your company") + ". With a strong background in software development and technical problem-solving as outlined in my resume, I am confident in my ability to contribute value to your team.\n\nThank you for your consideration.\n\nSincerely,\nCandidate"}