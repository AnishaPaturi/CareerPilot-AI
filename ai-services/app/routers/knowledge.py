from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from pydantic import BaseModel
from typing import List, Optional
import os
import json
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from datetime import datetime
from app.core.config import settings

router = APIRouter()

class ChatRequest(BaseModel):
    query: str
    document_ids: Optional[List[int]] = None

class ChatResponse(BaseModel):
    answer: str

class SummaryResponse(BaseModel):
    summary: str

class DocumentResponse(BaseModel):
    id: int
    filename: str
    upload_date: str
    chunk_count: int

# In-memory stores
VECTOR_STORES = {}
DOCUMENTS_STORE = {}
QUERY_HISTORY_STORE = {}
NOTES_STORE = {}
HIGHLIGHTS_STORE = {}
DOC_COUNTER = 0

def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        encode_kwargs={"normalize_embeddings": True}
    )

def get_vector_store(user_id: int):
    if user_id not in VECTOR_STORES:
        user_dir = f"./chroma_db/user_{user_id}"
        if os.path.exists(user_dir):
            VECTOR_STORES[user_id] = Chroma(
                persist_directory=user_dir,
                embedding_function=get_embeddings()
            )
        else:
            VECTOR_STORES[user_id] = None
    return VECTOR_STORES[user_id]

@router.post("/upload")
async def upload_document(file: UploadFile = File(...), user_id: int = 1):
    global DOC_COUNTER
    DOC_COUNTER += 1
    doc_id = DOC_COUNTER
    
    os.makedirs(f"./uploads/user_{user_id}", exist_ok=True)
    os.makedirs(f"./chroma_db/user_{user_id}", exist_ok=True)
    
    file_path = f"./uploads/user_{user_id}/{file.filename}"
    contents = await file.read()
    
    with open(file_path, "wb") as f:
        f.write(contents)
    
    loader = PyPDFLoader(file_path)
    documents = loader.load()
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)
    
    # Filter and clean chunks to ensure valid string content
    cleaned_chunks = []
    for idx, c in enumerate(chunks):
        content = c.page_content
        # Handle list or other non-string types
        if isinstance(content, list):
            content = " ".join(str(x) for x in content if x)
        elif not isinstance(content, str):
            content = str(content) if content else None
        if content:
            content = content.strip()
            if content:
                c.page_content = content
                c.metadata["chunk_index"] = idx
                cleaned_chunks.append(c)
    chunks = cleaned_chunks
    
    if not chunks:
        return {"id": doc_id, "filename": file.filename, "chunks": 0, "status": "no_valid_content", "message": "No valid text extracted from PDF"}
    
    vector_store = get_vector_store(user_id)
    if vector_store is not None:
        vector_store.add_documents(chunks)
    else:
        vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=get_embeddings(),
            persist_directory=f"./chroma_db/user_{user_id}"
        )
    
    VECTOR_STORES[user_id] = vector_store
    
    DOCUMENTS_STORE[doc_id] = {
        "id": doc_id,
        "user_id": user_id,
        "filename": file.filename,
        "upload_date": datetime.utcnow().isoformat(),
        "chunk_count": len(chunks)
    }
    
    if user_id not in QUERY_HISTORY_STORE:
        QUERY_HISTORY_STORE[user_id] = []
    
    return {"id": doc_id, "filename": file.filename, "chunks": len(chunks), "status": "success"}

@router.delete("/document/{doc_id}")
async def delete_document(doc_id: int, user_id: int = 1):
    if doc_id not in DOCUMENTS_STORE:
        raise HTTPException(status_code=404, detail="Document not found")
    doc = DOCUMENTS_STORE[doc_id]
    if doc["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    file_path = f"./uploads/user_{user_id}/{doc['filename']}"
    if os.path.exists(file_path):
        os.remove(file_path)
    
    del DOCUMENTS_STORE[doc_id]
    
    if user_id in VECTOR_STORES:
        VECTOR_STORES[user_id] = None
        user_dir = f"./chroma_db/user_{user_id}"
        if os.path.exists(user_dir):
            import shutil
            shutil.rmtree(user_dir)
    
    return {"message": "Document deleted"}

@router.get("/documents")
async def get_documents(user_id: int = 1):
    docs = [d for d in DOCUMENTS_STORE.values() if d["user_id"] == user_id]
    return {"documents": docs}

@router.get("/document/{doc_id}/file")
async def get_document_file(doc_id: int, user_id: int = 1):
    if doc_id not in DOCUMENTS_STORE:
        raise HTTPException(status_code=404, detail="Document not found")
    doc = DOCUMENTS_STORE[doc_id]
    if doc["user_id"] != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    file_path = f"./uploads/user_{user_id}/{doc['filename']}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    import urllib.parse
    encoded_filename = urllib.parse.quote(doc['filename'])
    return {"file_url": f"/uploads/user_{user_id}/{encoded_filename}"}

@router.post("/chat", response_model=ChatResponse)
async def chat_with_documents(request: ChatRequest, user_id: int = 1):
    vector_store = get_vector_store(user_id)
    
    if vector_store is None:
        raise HTTPException(status_code=400, detail="No documents uploaded yet")
    
    system_prompt = """You are a helpful AI assistant that answers questions about uploaded documents.
    Use the provided document context as the primary source.
    If the document lacks relevant info, use general knowledge while noting this.
    
    Context: {context}"""
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "{input}")
    ])
    
    retriever = vector_store.as_retriever(search_kwargs={"k": 10})
    
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    models = []
    for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
        if m and m not in models:
            models.append(m)
            
    last_error = None
    for model in models:
        try:
            llm = ChatOpenAI(
                api_key=os.getenv("OPENROUTER_API_KEY"),
                base_url="https://openrouter.ai/api/v1",
                model=model,
                temperature=0.1,
                max_retries=1
            )
            rag_chain = (
                {"context": retriever | format_docs, "input": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser()
            )
            answer = rag_chain.invoke(request.query)
            
            # Store query history
            query_entry = {
                "id": len(QUERY_HISTORY_STORE.get(user_id, [])) + 1,
                "query": request.query,
                "answer": answer,
                "timestamp": datetime.utcnow().isoformat()
            }
            if user_id not in QUERY_HISTORY_STORE:
                QUERY_HISTORY_STORE[user_id] = []
            QUERY_HISTORY_STORE[user_id].append(query_entry)
            
            return ChatResponse(answer=answer)
        except Exception as e:
            last_error = e
            import sys
            print(f"Failed chat_with_documents using model {model}: {e}", file=sys.stderr)
            continue
            
    raise HTTPException(status_code=500, detail=f"Chat failed. All models failed. Last error: {str(last_error)}")

@router.get("/history")
async def get_history(user_id: int = 1):
    return {"history": QUERY_HISTORY_STORE.get(user_id, [])}

@router.post("/summarize", response_model=SummaryResponse)
async def summarize_document(user_id: int = 1):
    try:
        vector_store = get_vector_store(user_id)
        
        if vector_store is None:
            raise HTTPException(status_code=400, detail="No documents uploaded")
        
        # Retrieve all documents to reconstruct full context in correct reading order
        try:
            collection = vector_store.get()
            documents_text = collection.get('documents', [])
            metadatas = collection.get('metadatas', [])
            
            # Group and sort chunks by document filename and page/index
            from collections import defaultdict
            grouped_docs = defaultdict(list)
            
            for idx, (doc_text, meta) in enumerate(zip(documents_text, metadatas)):
                source = meta.get('source', 'Unknown Document')
                filename = os.path.basename(source)
                page = meta.get('page', 0)
                chunk_idx = meta.get('chunk_index', idx)
                grouped_docs[filename].append((page, chunk_idx, doc_text))
            
            formatted_contexts = []
            for filename, chunks_list in grouped_docs.items():
                # Sort chunks by page first, then by chunk_index
                chunks_list.sort(key=lambda x: (x[0], x[1]))
                doc_context = "\n\n".join(text for page, chunk_idx, text in chunks_list)
                formatted_contexts.append(f"--- DOCUMENT: {filename} ---\n{doc_context}")
            
            context = "\n\n\n".join(formatted_contexts)
        except Exception as e:
            # Fallback to similarity search if .get() fails
            docs = vector_store.as_retriever(search_kwargs={"k": 50}).invoke("summarize main points")
            context = "\n\n".join(doc.page_content for doc in docs)
        
        system_prompt = (
            "You are an expert document analysis and synthesis AI. Your task is to generate a highly accurate, structured, "
            "and comprehensive summary of the provided document content.\n\n"
            "Strict Instructions:\n"
            "1. Grounding: Rely ONLY on the clear facts directly mentioned in the context. Do not assume, extrapolate, "
            "or use external/general knowledge. If a topic is not in the text, do not mention it.\n"
            "2. Structure: Format the summary beautifully with the following sections:\n"
            "   - **Executive Summary**: A concise, high-level overview of the document's purpose, main themes, and key conclusion.\n"
            "   - **Key Concepts & Core Topics**: A structured, bulleted list of the main concepts, subjects, or modules covered in the document.\n"
            "   - **Detailed Breakdown**: A detailed summary organized by sections, chapters, or topics as they appear in the source. "
            "Provide detailed summaries of findings, arguments, formulas, or methodologies presented in each section.\n"
            "   - **Actionable Takeaways & Key Takeouts**: Important takeaways, recommendations, conclusions, or highlights mentioned in the document.\n"
            "3. Multi-Document Handling: If the context contains multiple separate documents (indicated by '--- DOCUMENT: [filename] ---'), "
            "provide a distinct summary section for each document, clearly labeled with its filename, and then a final section comparing or connecting them if relevant.\n"
            "4. Tone and Style: Use professional, objective language. Keep formatting clean with clear markdown headings, bold text, and bullet points. "
            "Avoid generic summaries; ensure you mention specific terms, names, figures, and facts from the text.\n\n"
            "Document Content:\n"
            "{context}"
        )
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "Provide a detailed and highly accurate summary of the document(s) following the structural and grounding guidelines.")
        ])
        
        models = []
        for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
            if m and m not in models:
                models.append(m)
                
        last_error = None
        for model in models:
            try:
                llm = ChatOpenAI(
                    api_key=os.getenv("OPENROUTER_API_KEY"),
                    base_url="https://openrouter.ai/api/v1",
                    model=model,
                    temperature=0.1,
                    max_tokens=4096,
                    max_retries=1
                )
                summary = (prompt | llm | StrOutputParser()).invoke({"context": context})
                return SummaryResponse(summary=summary)
            except Exception as e:
                last_error = e
                import sys
                print(f"Failed summarize_document using model {model}: {e}", file=sys.stderr)
                continue
                
        raise HTTPException(
            status_code=500,
            detail=f"Summarization failed. All models failed. Last error: {str(last_error)}"
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/summary/convert/pdf")
async def convert_summary_to_pdf(summary: str):
    try:
        from weasyprint import HTML
        html_content = f"""
        <html>
        <head><style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #333; }}
            .summary {{ line-height: 1.6; }}
        </style></head>
        <body>
            <h1>Document Summary</h1>
            <div class="summary">{summary.replace('\n', '<br>')}</div>
        </body>
        </html>
        """
        pdf_bytes = HTML(string=html_content).write_pdf()
        import base64
        return {"pdf_base64": base64.b64encode(pdf_bytes).decode()}
    except ImportError:
        return {"error": "WeasyPrint not installed. Install with: pip install weasyprint"}

@router.post("/summary/convert/word")
async def convert_summary_to_word(summary: str):
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Document Summary', 0)
        doc.add_paragraph(summary)
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        import base64
        return {"word_base64": base64.b64encode(buffer.getvalue()).decode()}
    except ImportError:
        return {"error": "python-docx not installed. Install with: pip install python-docx"}

@router.post("/quiz")
async def generate_quiz(topic: str, user_id: int = 1, num_questions: int = 10):
    try:
        vector_store = get_vector_store(user_id)
        
        if vector_store is None:
            raise HTTPException(status_code=400, detail="No documents uploaded")
        
        docs = vector_store.as_retriever(search_kwargs={"k": 20}).invoke(topic)
        context = "\n\n".join(doc.page_content for doc in docs)
        
        prompt = ChatPromptTemplate.from_template("""
        Generate {num_questions} multiple choice questions about: {topic}
        Based on content: {context}
        
        Format JSON array: [{{"question", "options":["A","B","C","D"], "correct", "explanation"}}]
        """)
        
        models = []
        for m in [settings.OPENROUTER_MODEL, "google/gemma-4-31b-it:free", "google/gemma-4-26b-a4b-it:free", "poolside/laguna-m.1:free", "liquid/lfm-2.5-1.2b-thinking:free", "meta-llama/llama-3.2-3b-instruct:free", "openrouter/free"]:
            if m and m not in models:
                models.append(m)
                
        last_error = None
        for model in models:
            try:
                llm = ChatOpenAI(
                    api_key=os.getenv("OPENROUTER_API_KEY"),
                    base_url="https://openrouter.ai/api/v1",
                    model=model,
                    temperature=0.5,
                    max_retries=1
                )
                result = (prompt | llm | StrOutputParser()).invoke({"topic": topic, "num_questions": num_questions, "context": context})
                try:
                    parsed = json.loads(result)
                    return parsed
                except Exception as e:
                    # Clean markdown code blocks if any
                    cleaned = result.strip()
                    if cleaned.startswith("```json"):
                        cleaned = cleaned[7:]
                    if cleaned.endswith("```"):
                        cleaned = cleaned[:-3]
                    cleaned = cleaned.strip()
                    parsed = json.loads(cleaned)
                    return parsed
            except Exception as e:
                last_error = e
                import sys
                print(f"Failed generate_quiz using model {model}: {e}", file=sys.stderr)
                continue
                
        raise HTTPException(
            status_code=500,
            detail=f"Quiz generation failed. All models failed. Last error: {str(last_error)}"
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class NoteCreate(BaseModel):
    document_id: int = 1
    content: str
    page_number: int = 1

class NoteResponse(BaseModel):
    id: int
    document_id: int
    content: str
    page_number: int
    created_at: str

@router.post("/notes", response_model=NoteResponse)
async def create_note(note: NoteCreate, user_id: int = 1):
    new_note = {
        "id": len(NOTES_STORE.get(user_id, [])) + 1,
        "user_id": user_id,
        "document_id": note.document_id,
        "content": note.content,
        "page_number": note.page_number,
        "created_at": datetime.utcnow().isoformat()
    }
    if user_id not in NOTES_STORE:
        NOTES_STORE[user_id] = []
    NOTES_STORE[user_id].append(new_note)
    return new_note

@router.get("/notes")
async def get_notes(user_id: int = 1):
    notes = NOTES_STORE.get(user_id, [])
    return {"notes": notes}

@router.delete("/notes/{note_id}")
async def delete_note(note_id: int, user_id: int = 1):
    if user_id in NOTES_STORE:
        NOTES_STORE[user_id] = [n for n in NOTES_STORE[user_id] if n["id"] != note_id]
    return {"message": "Note deleted"}

class HighlightCreate(BaseModel):
    document_id: int = 1
    page_number: int = 1
    text: str
    x: float = 0
    y: float = 0

class HighlightResponse(BaseModel):
    id: int
    document_id: int
    page_number: int
    text: str
    x: float
    y: float

@router.post("/highlights", response_model=HighlightResponse)
async def create_highlight(highlight: HighlightCreate, user_id: int = 1):
    new_highlight = {
        "id": len(HIGHLIGHTS_STORE.get(user_id, [])) + 1,
        "user_id": user_id,
        "document_id": highlight.document_id,
        "page_number": highlight.page_number,
        "text": highlight.text,
        "x": highlight.x,
        "y": highlight.y
    }
    if user_id not in HIGHLIGHTS_STORE:
        HIGHLIGHTS_STORE[user_id] = []
    HIGHLIGHTS_STORE[user_id].append(new_highlight)
    return new_highlight

@router.get("/highlights")
async def get_highlights(user_id: int = 1):
    highlights = HIGHLIGHTS_STORE.get(user_id, [])
    return {"highlights": highlights}

@router.delete("/highlights/{highlight_id}")
async def delete_highlight(highlight_id: int, user_id: int = 1):
    if user_id in HIGHLIGHTS_STORE:
        HIGHLIGHTS_STORE[user_id] = [h for h in HIGHLIGHTS_STORE[user_id] if h["id"] != highlight_id]
    return {"message": "Highlight deleted"}

@router.post("/chat/quote")
async def chat_with_quote(request: ChatRequest, user_id: int = 1):
    return await chat_with_documents(request, user_id)